# backend/utils/file_parser.py
import os
import base64
import fitz  # PyMuPDF，用于处理 PDF
import docx  # python-docx，用于处理 Word

class FileProcessor:
    def __init__(self, is_multimodal=False):
        # 这个开关决定了系统现在是纯文本还是多模态
        # 现阶段（纯文本模型）设为 False，走传统文字提取。
        # 后续换成 GLM-4V 等多模态模型时，只需在这里改为 True 即可！
        self.is_multimodal = is_multimodal

    def process_file(self, file_path: str):
        """核心入口：根据文件类型和模型能力，返回大模型能接受的 list 数据结构"""
        if not os.path.exists(file_path):
            return None

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._handle_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._handle_docx(file_path)
        elif ext in [".jpg", ".jpeg", ".png", ".webp"]:
            return self._handle_image(file_path)
        else:
            return [{"type": "text", "text": f"\n[系统提示：不支持解析该文件类型 ({ext})]"}]

    def _handle_pdf(self, file_path):
        """处理 PDF：文本模型抽文字，多模态模型直接截高清图"""
        try:
            doc = fitz.open(file_path)
            
            if not self.is_multimodal:
                # 当前模式：纯文本模型。剥离出文字，丢弃图片。
                text = ""
                # 限制只读取前 20 页，保护 tokens 不超载
                for page_num in range(min(20, doc.page_count)):
                    page = doc.load_page(page_num)
                    text += page.get_text()
                    
                return [{
                    "type": "text", 
                    "text": f"\n--- 以下是用户上传的 PDF 文档纯文本内容 ---\n{text}\n--- PDF 文档结束 ---\n"
                }]
                
            else:
                # 未来模式：多模态模型！把每一页变成高清图片，让模型直接“看”！
                content_list = [{
                    "type": "text",
                    "text": "以下是用户上传的 PDF 文档的页面截图，请结合图片内容回答问题："
                }]
                
                # 逐页转图片 (这里限制前 5 页，因为图片极其消耗 Token)
                for page_num in range(min(5, doc.page_count)):
                    page = doc.load_page(page_num)
                    # 放大 2 倍，让生成的图片更清晰，防止大模型看不清小字
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    
                    img_bytes = pix.tobytes("png")
                    encoded_string = base64.b64encode(img_bytes).decode('utf-8')
                    
                    content_list.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{encoded_string}"
                        }
                    })
                return content_list

        except Exception as e:
            return [{"type": "text", "text": f"\n[系统提示：PDF 解析失败 ({str(e)})]"}]

    def _handle_docx(self, file_path):
        """处理 Word：提取文本、表格，并对文档内的图片进行多模态/OCR处理"""
        try:
            doc = docx.Document(file_path)
            full_text = []

            # 1. 提取所有段落文字
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # 2. 提取并转换所有表格为 Markdown
            if doc.tables:
                full_text.append("\n--- 文档中的表格数据 ---")
                for table in doc.tables:
                    for i, row in enumerate(table.rows):
                        row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        full_text.append("| " + " | ".join(row_data) + " |")
                        if i == 0:
                            full_text.append("|" + "|".join(["---"] * len(row.cells)) + "|")
                    full_text.append("\n") 

            final_text_content = "\n".join(full_text)
            
            # 基础文本结构
            content_list = [{
                "type": "text",
                "text": f"\n--- 以下是用户上传的 Word 文档文本和表格内容 ---\n{final_text_content}\n--- Word 文本结束 ---\n"
            }]

            # 处理 Word 里的图片！
            # 遍历 Word 文档底层的关系(rels)，找出所有图片部件
            image_ocr_texts = []
            
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_bytes = rel.target_part.blob  # 获取图片的真实二进制流
                    
                    if self.is_multimodal:
                        # 多模态模式：直接把 Word 里的图片转成 Base64 塞给大模型
                        encoded_string = base64.b64encode(image_bytes).decode('utf-8')
                        content_list.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{encoded_string}" 
                            }
                        })
                    else:
                        # 纯文本模式：用 OCR 把图片里的字抠出来
                        if self.ocr_reader:
                            ocr_results = self.ocr_reader.readtext(image_bytes, detail=0)
                            if ocr_results:
                                image_ocr_texts.append(" ".join(ocr_results))

            # 如果是纯文本模式，且在图片中提取到了 OCR 文字，把它追加给大模型
            if not self.is_multimodal and image_ocr_texts:
                ocr_combined_text = "\n".join([f"插图 {i+1}：{t}" for i, t in enumerate(image_ocr_texts)])
                content_list.append({
                    "type": "text",
                    "text": f"\n--- 以下是 Word 文档插图的 OCR 扫描结果 ---\n{ocr_combined_text}\n"
                })

            return content_list

        except Exception as e:
            return [{"type": "text", "text": f"\n[系统提示：Word 解析失败 ({str(e)})]"}]

    def _handle_image(self, file_path):
        """处理图片：根据模型能力决定是否进行 Base64 编码"""
        if self.is_multimodal:
            # 模型能看图！把图片转成 Base64 喂给它
            try:
                with open(file_path, "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                    return [{
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/jpeg;base64,{encoded_string}"
                        }
                    }]
            except Exception as e:
                return [{"type": "text", "text": f"\n[系统提示：图片解析失败 ({str(e)})]"}]
        else:
            # 当前模式：纯文本模型，瞎子。只能给个文字提示。
            filename = os.path.basename(file_path)
            return [{
                "type": "text", 
                "text": f"\n[系统提示：用户向你发送了一张名为 '{filename}' 的图片，但你是一个纯文本模型，当前无法查看图片内容。请礼貌地告知用户你看不见图片。]\n"
            }]